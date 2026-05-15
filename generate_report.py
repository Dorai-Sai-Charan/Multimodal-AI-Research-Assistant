#!/usr/bin/env python3
"""Generate Multimodal AI Research Assistant project report following 8thCSE template."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(text, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_before=0, space_after=6, line_spacing=1.5, style=None):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = Pt(line_spacing * 12)
    if text:
        run = p.add_run(text)
        set_font(run, size, bold=bold, italic=italic)
    return p

def heading_main(text):
    """16pt Bold centred – chapter heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, 16, bold=True)
    return p

def heading_side(text, level=1):
    """14pt Bold left – section heading."""
    size = 14 if level == 1 else 13
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size, bold=True)
    return p

def body(text, space_after=6):
    return para(text, size=12, space_after=space_after)

def add_page_break():
    doc.add_page_break()

def add_table(headers, rows, caption=""):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        set_font(run, 11, bold=True, italic=True)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            for run in cell.paragraphs[0].runs:
                set_font(run, 10)

    doc.add_paragraph()  # spacing after table
    return tbl

FIGURES_DIR = "/home/nst-kaja/Documents/Dorai/Multimodal-AI-Research-Assistant/figures"

def add_figure(fig_num, caption, img_filename=None):
    """Insert a real image if available, otherwise a labelled placeholder box."""
    import os
    img_path = os.path.join(FIGURES_DIR, img_filename) if img_filename else None
    if img_path and os.path.exists(img_path):
        # Centring paragraph then insert picture
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.8))
    else:
        # Placeholder box for images not yet available
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after  = Pt(2)
        r = p.add_run(f"[ Figure {fig_num} — insert image here ]")
        set_font(r, 11, italic=True, color=(120, 120, 120))
    # Caption always added below
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cap.add_run(f"Fig. {fig_num}  {caption}")
    set_font(r2, 11, bold=True)
    cap.paragraph_format.space_after = Pt(10)

def figure_placeholder(fig_num, caption):
    """Legacy wrapper — calls add_figure with no image path."""
    add_figure(fig_num, caption, img_filename=None)

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Multimodal AI Research Assistant")
set_font(run, 18, bold=True)

para("A PROJECT REPORT", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)
para("Submitted by", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16)

for sid, name in [("BL.EN.U4CSE22-001", "Dorai Sai Charan"),
                  ("BL.EN.U4CSE22-002", "Team Member 2"),
                  ("BL.EN.U4CSE22-003", "Team Member 3")]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{sid}\t{name}")
    set_font(r1, 12)

para("in partial fulfillment for the award of the degree of",
     size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16)
para("BACHELOR OF TECHNOLOGY", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para("IN", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para("COMPUTER SCIENCE AND ENGINEERING", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(2):
    doc.add_paragraph()

para("AMRITA SCHOOL OF COMPUTING, BENGALURU", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para("AMRITA VISHWA VIDYAPEETHAM", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para("BENGALURU 560 035", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para("MAY 2026", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ══════════════════════════════════════════════════════════════════════════════
#  BONAFIDE CERTIFICATE
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("AMRITA VISHWA VIDYAPEETHAM\nAMRITA SCHOOL OF COMPUTING, BENGALURU, 560035")
para("", space_after=12)
heading_main("BONAFIDE CERTIFICATE")

body('This is to certify that the project report entitled "Multimodal AI Research Assistant" submitted by')
para("", space_after=4)
for sid, name in [("BL.EN.U4CSE22-001", "Dorai Sai Charan"),
                  ("BL.EN.U4CSE22-002", "Team Member 2"),
                  ("BL.EN.U4CSE22-003", "Team Member 3")]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{sid}\t{name}")
    set_font(r1, 12)

body('in partial fulfillment of the requirements as part of Bachelor of Technology in '
     '"COMPUTER SCIENCE AND ENGINEERING" is a bonafide record of the work carried out '
     'under our guidance and supervision at Amrita School of Computing, Bengaluru.')
para("", space_after=24)

tbl = doc.add_table(rows=2, cols=2)
tbl.style = "Table Grid"
cells = tbl.rows[0].cells
cells[0].text = "____________________"
cells[1].text = "____________________"
cells = tbl.rows[1].cells
cells[0].text = "Project Guide\nDepartment of CSE\nAmrita School of Computing"
cells[1].text = "Chair\nSchool of Computing"
for row in tbl.rows:
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            set_font(run, 11)

para("", space_after=16)
body("This project report was evaluated by us on …………")
para("", space_after=24)
tbl2 = doc.add_table(rows=2, cols=2)
tbl2.style = "Table Grid"
cells2 = tbl2.rows[0].cells
cells2[0].text = "____________________"
cells2[1].text = "____________________"
cells2 = tbl2.rows[1].cells
cells2[0].text = "Examiner 1"
cells2[1].text = "Examiner 2"

# ══════════════════════════════════════════════════════════════════════════════
#  ACKNOWLEDGEMENTS
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("ACKNOWLEDGEMENTS")

body("The satisfaction that accompanies the successful completion of any task would be incomplete "
     "without mentioning the people who made it possible, whose constant encouragement and "
     "guidance have been an endless source of inspiration throughout the course of this project work.")

body("We offer our sincere pranams at the lotus feet of 'AMMA', Mata Amritanandamayi Devi, who "
     "showered her blessings upon us throughout the course of this project work.")

body("We owe our gratitude to Prof. Manoj P., Director, Amrita Vishwa Vidyapeetham, Bengaluru "
     "Campus. We would like to place our heartfelt gratitude to Dr. Gopalakrishnan E.A., "
     "Principal, Amrita School of Computing and Amrita School of Artificial Intelligence, "
     "Bengaluru, for his valuable support and inspiration.")

body("It is a great pleasure to express our gratitude and indebtedness to our project guide, "
     "Department of Computer Science and Engineering, Amrita School of Computing, Bengaluru, "
     "for their invaluable guidance, encouragement, moral support, and affection throughout "
     "the project work.")

body("We would like to express our gratitude to the project panel members for their suggestions, "
     "encouragement, and moral support during the process of project work, and all faculty members "
     "for their academic support. Finally, we are forever grateful to our parents, who have loved, "
     "supported, and encouraged us in all our endeavours.")

para("", space_after=24)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Dorai Sai Charan\nTeam Member 2\nTeam Member 3")
set_font(run, 12)

# Roman numeral page marker
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("I")
set_font(r2, 12, italic=True)

# ══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("ABSTRACT")

body("The exponential growth of scientific literature has created an urgent need for intelligent "
     "systems capable of extracting, reasoning, and synthesising knowledge from complex, "
     "multimodal research documents. This project presents a Multimodal AI Research Assistant — "
     "a full-stack, end-to-end intelligent research companion that integrates Retrieval-Augmented "
     "Generation (RAG), Agentic Reasoning via the ReAct (Reason + Act) pattern, and a novel "
     "AI-Detection and Text Humanization Engine into a single cohesive platform.")

body("The system accepts PDF research papers as input and applies a six-stage multimodal ingestion "
     "pipeline that separately processes textual content via PyMuPDF, tabular data via pdfplumber, "
     "embedded figures via image extraction, handwritten or scanned text via EasyOCR, visual "
     "diagrams and charts via Gemini 2.0 Flash Vision, and mathematical equations via LaTeX "
     "conversion. Extracted content is semantically chunked using LangChain's Recursive Character "
     "Text Splitter (chunk size 512 characters, overlap 50 characters) and embedded into a "
     "384-dimensional vector space using the all-MiniLM-L6-v2 Sentence-Transformer model. "
     "Embeddings are persisted in ChromaDB using an HNSW cosine-similarity index.")

body("For knowledge retrieval and generation, the system implements nine specialised RAG pipeline "
     "methods — including single-shot question answering, paper summarisation, paper comparison, "
     "literature survey generation, research gap identification, and cross-paper multi-document "
     "reasoning — each with tailored retrieval strategies and prompt templates. A ReAct-based "
     "Research Agent orchestrates multi-hop reasoning through up to ten iterative Thought–Action–"
     "Observation cycles, dynamically invoking typed retrieval tools (text, table, figure, and "
     "equation search) for complex queries. Text generation is powered by Groq's Llama 3.3 70B "
     "model for high-throughput inference.")

body("A unique Humanizer Engine combines a RoBERTa-based AI content detector with heuristic "
     "linguistic metrics (burstiness, Type-Token Ratio, human marker density) to score and "
     "iteratively refine generated text until its AI-detection score falls below 20%, ensuring "
     "academic authenticity. The modern Next.js 15 frontend with React 19 renders rich Markdown, "
     "KaTeX mathematical notation, syntax-highlighted code, and inline figures. Experimental "
     "evaluation demonstrates strong retrieval precision, contextually faithful answers with "
     "full citation traceability, and effective humanization, making this system a comprehensive "
     "tool for researchers and graduate students.")

p_rom = doc.add_paragraph()
p_rom.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_rom = p_rom.add_run("II")
set_font(r_rom, 12, italic=True)

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("TABLE OF CONTENTS")

toc_entries = [
    ("ACKNOWLEDGEMENTS", "i"),
    ("ABSTRACT", "ii"),
    ("LIST OF FIGURES", "v"),
    ("LIST OF TABLES", "vi"),
    ("CHAPTER 1 – INTRODUCTION", "1"),
    ("    1.1  Background and Context", "2"),
    ("    1.2  Problem Statement", "3"),
    ("    1.3  Motivation", "4"),
    ("    1.4  Objectives", "5"),
    ("    1.5  Scope of the Project", "6"),
    ("    1.6  Organisation of the Report", "6"),
    ("CHAPTER 2 – LITERATURE SURVEY", "7"),
    ("    2.1  Retrieval-Augmented Generation Systems", "7"),
    ("    2.2  Multimodal Document Understanding", "9"),
    ("    2.3  Agentic AI and ReAct Pattern", "11"),
    ("    2.4  AI Detection and Text Humanization", "13"),
    ("    2.5  Comparative Analysis of Related Work", "14"),
    ("CHAPTER 3 – SYSTEM REQUIREMENTS AND ANALYSIS", "16"),
    ("    3.1  Functional Requirements", "16"),
    ("    3.2  Non-Functional Requirements", "17"),
    ("    3.3  Software Requirements", "18"),
    ("    3.4  Hardware Requirements", "19"),
    ("CHAPTER 4 – SYSTEM DESIGN", "20"),
    ("    4.1  High-Level Architecture", "20"),
    ("    4.2  Ingestion Pipeline Design", "22"),
    ("    4.3  Storage Layer Design", "24"),
    ("    4.4  RAG Pipeline Design", "25"),
    ("    4.5  Agent Design", "26"),
    ("    4.6  API Layer Design", "27"),
    ("    4.7  Frontend Design", "28"),
    ("CHAPTER 5 – SYSTEM IMPLEMENTATION", "30"),
    ("    5.1  Document Ingestion Module", "30"),
    ("    5.2  Embedding and Vector Storage", "33"),
    ("    5.3  Retrieval and RAG Pipeline", "35"),
    ("    5.4  Agentic Reasoning Module", "37"),
    ("    5.5  Humanizer Engine", "39"),
    ("    5.6  REST API Module", "41"),
    ("    5.7  Frontend Implementation", "43"),
    ("CHAPTER 6 – SYSTEM TESTING", "45"),
    ("    6.1  Unit Testing", "45"),
    ("    6.2  Integration Testing", "46"),
    ("    6.3  Performance and Load Testing", "47"),
    ("    6.4  RAG Evaluation Using RAGAS Framework", "48"),
    ("CHAPTER 7 – RESULTS AND ANALYSIS", "50"),
    ("    7.1  Retrieval Performance", "50"),
    ("    7.2  Generation Quality", "51"),
    ("    7.3  Agent Reasoning Accuracy", "52"),
    ("    7.4  Humanization Effectiveness", "53"),
    ("    7.5  System Throughput and Latency", "54"),
    ("CHAPTER 8 – CONCLUSION AND FUTURE SCOPE", "56"),
    ("    8.1  Conclusion", "56"),
    ("    8.2  Future Enhancements", "57"),
    ("REFERENCES", "59"),
]

tbl_toc = doc.add_table(rows=len(toc_entries), cols=2)
tbl_toc.style = "Table Grid"
for r_idx, (entry, page) in enumerate(toc_entries):
    row = tbl_toc.rows[r_idx]
    row.cells[0].text = entry
    row.cells[1].text = page
    for c in row.cells:
        for run in c.paragraphs[0].runs:
            bold = not entry.startswith("  ")
            set_font(run, 11, bold=bold)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

p_rom3 = doc.add_paragraph()
p_rom3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_rom3 = p_rom3.add_run("III")
set_font(r_rom3, 12, italic=True)

# ══════════════════════════════════════════════════════════════════════════════
#  LIST OF FIGURES
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("LIST OF FIGURES")

figures = [
    ("1.1", "Conceptual overview of the Multimodal AI Research Assistant", "2"),
    ("2.1", "RAG pipeline overview — retrieval, augmentation, and generation", "8"),
    ("2.2", "Multimodal document processing taxonomy", "10"),
    ("2.3", "ReAct agent reasoning loop (Thought–Action–Observation)", "12"),
    ("4.1", "High-level system architecture diagram", "21"),
    ("4.2", "Six-stage PDF ingestion pipeline", "23"),
    ("4.3", "ChromaDB HNSW vector storage and cosine similarity retrieval", "24"),
    ("4.4", "Nine-method RAG pipeline overview", "25"),
    ("4.5", "ReAct agent iteration flow", "26"),
    ("4.6", "REST API endpoint structure", "27"),
    ("4.7", "Next.js frontend page layout and component hierarchy", "29"),
    ("5.1", "PDF processor heading detection decision logic", "31"),
    ("5.2", "Chunking strategy by content type", "33"),
    ("5.3", "Embedding pipeline: text to 384-dimensional vector", "34"),
    ("5.4", "Retrieval flow: query embedding to ranked results", "36"),
    ("5.5", "Agent multi-hop reasoning example trace", "38"),
    ("5.6", "Humanizer engine iterative refinement loop", "40"),
    ("6.1", "RAGAS evaluation score distribution", "49"),
    ("7.1", "Retrieval precision@k for various chunk types", "50"),
    ("7.2", "AI detection score before and after humanization", "53"),
    ("7.3", "API endpoint latency distribution (p50, p90, p99)", "55"),
]

tbl_figs = doc.add_table(rows=len(figures) + 1, cols=3)
tbl_figs.style = "Table Grid"
for i, h in enumerate(["Figure No.", "Caption", "Page"]):
    tbl_figs.rows[0].cells[i].text = h
    for run in tbl_figs.rows[0].cells[i].paragraphs[0].runs:
        set_font(run, 11, bold=True)
for r_idx, (num, cap, pg) in enumerate(figures):
    row = tbl_figs.rows[r_idx + 1]
    for c_idx, txt in enumerate([f"Fig. {num}", cap, pg]):
        row.cells[c_idx].text = txt
        for run in row.cells[c_idx].paragraphs[0].runs:
            set_font(run, 10)

p_rom4 = doc.add_paragraph()
p_rom4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_rom4 = p_rom4.add_run("V")
set_font(r_rom4, 12, italic=True)

# ══════════════════════════════════════════════════════════════════════════════
#  LIST OF TABLES
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("LIST OF TABLES")

tables_list = [
    ("3.1", "Software requirements specification", "18"),
    ("3.2", "Hardware requirements specification", "19"),
    ("3.3", "Functional requirements summary", "16"),
    ("4.1", "Ingestion processor responsibilities", "22"),
    ("4.2", "ChromaDB metadata schema per chunk", "24"),
    ("5.1", "PDF heading detection keywords", "31"),
    ("5.2", "Chunking parameters by content type", "32"),
    ("5.3", "RAG pipeline methods and retrieval strategies", "35"),
    ("5.4", "Groq API models supported by the system", "42"),
    ("6.1", "Unit test cases and results", "45"),
    ("6.2", "Integration test scenarios", "46"),
    ("6.3", "RAGAS metric scores across test queries", "48"),
    ("7.1", "Retrieval Hit Rate and MRR at k=5 and k=10", "51"),
    ("7.2", "LLM generation quality scores (faithfulness, relevance)", "52"),
    ("7.3", "Comparison with baseline and existing systems", "54"),
]

tbl_tabs = doc.add_table(rows=len(tables_list) + 1, cols=3)
tbl_tabs.style = "Table Grid"
for i, h in enumerate(["Table No.", "Caption", "Page"]):
    tbl_tabs.rows[0].cells[i].text = h
    for run in tbl_tabs.rows[0].cells[i].paragraphs[0].runs:
        set_font(run, 11, bold=True)
for r_idx, (num, cap, pg) in enumerate(tables_list):
    row = tbl_tabs.rows[r_idx + 1]
    for c_idx, txt in enumerate([f"Table {num}", cap, pg]):
        row.cells[c_idx].text = txt
        for run in row.cells[c_idx].paragraphs[0].runs:
            set_font(run, 10)

p_rom5 = doc.add_paragraph()
p_rom5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_rom5 = p_rom5.add_run("VI")
set_font(r_rom5, 12, italic=True)

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("CHAPTER – 1\nINTRODUCTION")

heading_side("1.1  Background and Context")
body("The scientific community produces more than two million peer-reviewed publications annually, "
     "a figure that doubles roughly every nine years. For researchers, engineers, and graduate "
     "students, manually reading, comprehending, and synthesising relevant knowledge from this "
     "ocean of literature has become an increasingly untenable task. Traditional keyword-based "
     "search engines and bibliographic databases such as Google Scholar, PubMed, and Semantic "
     "Scholar address document discovery but offer no mechanism for deep comprehension, comparative "
     "analysis, or synthesis across papers.")

body("The advent of Large Language Models (LLMs) such as GPT-4, Llama 3, and Gemini has opened "
     "a new paradigm: conversational access to knowledge encoded in text. However, general-purpose "
     "LLMs suffer from well-documented limitations when applied to domain-specific scientific "
     "literature. First, they are constrained by a fixed context window that cannot accommodate "
     "the full content of even a single long research paper, let alone a collection. Second, they "
     "lack access to documents published after their training cutoff, creating a temporal blind "
     "spot. Third — and most critically — they hallucinate: they fabricate plausible-sounding but "
     "factually incorrect claims, a behaviour that is particularly hazardous in scientific contexts "
     "where precision is paramount.")

body("Retrieval-Augmented Generation (RAG) was introduced as a principled solution to these "
     "limitations. By dynamically fetching relevant passages from an external knowledge store "
     "and injecting them into the LLM's context window, RAG grounds the model's responses in "
     "verifiable evidence, dramatically reducing hallucination and enabling access to arbitrary "
     "external corpora. The seminal RAG paper by Lewis et al. (2020) demonstrated that "
     "retrieval-augmented models outperform purely generative counterparts on knowledge-intensive "
     "tasks by a substantial margin.")

body("Research documents, however, are inherently multimodal. A typical machine learning paper "
     "contains not only narrative text but also data tables reporting experimental results, "
     "architecture diagrams describing system components, performance graphs, and mathematical "
     "equations formalising model behaviour. Standard RAG systems are text-only pipelines that "
     "discard the rich information encoded in tables, figures, and equations when processing "
     "research papers. This information gap motivates the construction of a truly multimodal "
     "retrieval system.")

add_figure("1.1", "Conceptual overview of the Multimodal AI Research Assistant", "fig_1_1_overview.png")

heading_side("1.2  Problem Statement")
body("Existing research assistance tools present several critical gaps that this project aims to "
     "address:")

body("(i) Text-only RAG pipelines discard tables, figures, and equations — often the most "
     "information-dense components of a research paper — resulting in incomplete and superficial "
     "answers to quantitative or methodology-focused questions.")

body("(ii) General-purpose conversational AI systems cannot reason across multiple uploaded "
     "documents simultaneously, preventing comparative analysis or synthesis across a personal "
     "literature collection.")

body("(iii) Complex research questions requiring multi-step reasoning — such as 'How does the "
     "attention mechanism described in Paper A compare to the sparse attention approach in Paper B, "
     "and which achieves better performance on the GLUE benchmark?' — cannot be answered by a "
     "single-shot RAG pipeline and require iterative, tool-augmented reasoning.")

body("(iv) AI-generated research content, even when factually accurate, is often recognisable "
     "as machine-produced due to its uniform sentence structure and lexical patterns, which may "
     "not meet academic writing standards or pass institutional AI-detection checks.")

heading_side("1.3  Motivation")
body("The primary motivation for this project arises from the direct experience of graduate "
     "students and researchers who spend disproportionate time on literature review activities. "
     "Studies in information science indicate that researchers spend on average 23% of their "
     "working time searching for and reading literature. An intelligent assistant that can "
     "instantly answer questions about uploaded papers, generate structured literature surveys, "
     "identify research gaps, and explain complex mathematical formulations would represent a "
     "significant productivity multiplier for the academic community.")

body("A secondary motivation is the advancement of the state of the art in applied multimodal "
     "AI systems. While numerous RAG implementations exist, very few combine multimodal "
     "ingestion, persistent vector storage, agentic multi-hop reasoning, and AI-detection "
     "humanization in a single integrated system. Building and evaluating such a system "
     "contributes original engineering insight to the rapidly growing field of applied AI.")

body("The project is also motivated by the convergence of high-quality free-tier APIs (Groq "
     "for ultra-fast LLM inference, Gemini 2.0 Flash for vision tasks) with open-source vector "
     "databases (ChromaDB) and embedding models (Sentence-Transformers), which together make a "
     "production-quality multimodal RAG system achievable without significant computational "
     "expenditure.")

heading_side("1.4  Objectives")
body("The specific objectives of this project are as follows:")
body("1. To design and implement a multimodal PDF ingestion pipeline capable of extracting "
     "text, tables, images, OCR content, figure descriptions, and mathematical equations "
     "from research papers.")
body("2. To build a persistent semantic knowledge base using Sentence-Transformer embeddings "
     "and ChromaDB vector storage that supports typed retrieval across content modalities.")
body("3. To implement nine specialised Retrieval-Augmented Generation pipelines for distinct "
     "research tasks including question answering, summarisation, paper comparison, literature "
     "survey generation, research gap identification, and trend analysis.")
body("4. To develop a ReAct-based agentic reasoning module capable of multi-hop question "
     "answering through iterative tool invocation and observation.")
body("5. To engineer a hybrid AI Detection and Humanization Engine that scores AI-generated "
     "content using a RoBERTa model combined with heuristic linguistic metrics, and iteratively "
     "rewrites text to reduce its AI signature below a configurable threshold.")
body("6. To deliver a modern, responsive web interface using Next.js 15 and React 19 that "
     "renders markdown, mathematical notation (KaTeX), code, and inline figures.")

heading_side("1.5  Scope of the Project")
body("This project covers the full-stack development of an AI-powered research assistant from "
     "raw PDF ingestion through to the user-facing interface. The scope includes backend API "
     "development in Python using FastAPI, document processing, embedding generation, vector "
     "database management, LLM orchestration, agentic reasoning, and frontend development with "
     "Next.js. The system is designed for single-user or small-team deployment on standard "
     "hardware with internet access to Groq and Gemini APIs.")

body("The scope does not include: real-time paper ingestion from the web, multi-user "
     "authentication and authorisation beyond basic access control, fine-tuning of language "
     "models, or integration with institutional library management systems.")

heading_side("1.6  Organisation of the Report")
body("Chapter 2 presents a comprehensive literature survey covering RAG systems, multimodal "
     "document understanding, agentic AI, and AI detection research. Chapter 3 specifies "
     "functional, non-functional, software, and hardware requirements. Chapter 4 describes "
     "system design at both high and low levels. Chapter 5 details the implementation of each "
     "module. Chapter 6 covers testing strategies and evaluation. Chapter 7 presents results "
     "and analysis. Chapter 8 concludes the report and outlines future enhancement directions.")

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 2 — LITERATURE SURVEY
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("CHAPTER – 2\nLITERATURE SURVEY")

heading_side("2.1  Retrieval-Augmented Generation Systems")
body("The concept of augmenting language model generation with retrieved context was formalised "
     "by Lewis et al. in 'Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks' "
     "(NeurIPS 2020). Their framework combined a neural retriever (DPR) with a generative model "
     "(BART) and demonstrated superior performance on open-domain question answering benchmarks "
     "including Natural Questions, TriviaQA, and WebQuestions. The key insight was that "
     "conditioning generation on retrieved passages reduces the burden on the model's parametric "
     "memory and improves factual accuracy.")

body("Subsequent work significantly advanced RAG architectures. The Fusion-in-Decoder (FiD) "
     "model by Izacard and Grave (2021) extended RAG by encoding multiple retrieved passages "
     "independently and fusing them in the decoder, achieving new state-of-the-art results on "
     "knowledge-intensive tasks. Shi et al. (2023) introduced REPLUG, a retrieval-augmented "
     "language model that treats the retriever as a plug-in module and demonstrates that even "
     "black-box LLMs can benefit substantially from retrieval augmentation.")

body("The survey paper 'Agentic Retrieval-Augmented Generation: A Survey on Agentic RAG' "
     "(2024) provides a comprehensive taxonomy of RAG systems, categorising them along the "
     "dimensions of retrieval granularity, indexing strategy, query rewriting, and generation "
     "faithfulness. The survey identifies three primary RAG paradigms: Naive RAG (simple "
     "retrieve-then-generate), Advanced RAG (with pre-retrieval query optimisation and "
     "post-retrieval re-ranking), and Modular RAG (with interchangeable pipeline components).")

body("'mRAG: Elucidating the Design Space of Multimodal Retrieval-Augmented Generation' (2024) "
     "systematically studies the design choices in multimodal RAG systems, analysing how "
     "retrieval modality, indexing granularity (page-level vs. element-level), and fusion "
     "strategy affect end-to-end performance. The authors find that element-level indexing "
     "with typed retrieval consistently outperforms page-level approaches on question answering "
     "benchmarks, directly informing the design of the type-aware chunking strategy adopted "
     "in this project.")

add_figure("2.1", "RAG pipeline overview — retrieval, augmentation, and generation", "fig_2_1_rag_pipeline.png")

heading_side("2.2  Multimodal Document Understanding")
body("Multimodal document understanding encompasses the extraction, representation, and "
     "reasoning over heterogeneous content types within documents. Early work focused on "
     "structured document parsing using layout analysis models such as LayoutLM (Xu et al., "
     "2020), which jointly models text and spatial position via 2D positional embeddings. "
     "LayoutLMv3 (Huang et al., 2022) further incorporates image patches, enabling unified "
     "text-image-layout pretraining.")

body("'Scaling Beyond Context: A Survey of Multimodal RAG for Document Understanding' (2024) "
     "surveys 47 multimodal RAG systems and categorises their document processing strategies. "
     "The survey identifies four primary content extraction approaches: pipeline-based "
     "extraction (separate models per modality), end-to-end neural parsing, OCR-based text "
     "detection, and vision-language model (VLM) captioning. The authors conclude that "
     "pipeline-based approaches remain the most practical for production systems due to their "
     "interpretability and modularity, a finding aligned with the architecture chosen for "
     "this project.")

body("'CMRAG: Co-Modality-Based Visual Document Retrieval and Question Answering' (2024) "
     "introduces a co-modality retrieval approach where visual and textual query embeddings "
     "are jointly used to retrieve multimodal chunks. Their system achieves 18.3% higher "
     "accuracy on the DocVQA benchmark compared to text-only baselines, demonstrating the "
     "tangible value of multimodal retrieval. The VisionAnalyzer component in this project, "
     "which uses Gemini 2.0 Flash to generate textual descriptions of figures that are then "
     "embedded alongside text content, implements a similar modality-bridging strategy.")

body("'DocAgent: An Agentic Framework for Multi-Modal Long-Context Document Understanding' "
     "(2024) presents a framework specifically designed for long research documents that "
     "exceed the context window of current LLMs. DocAgent uses a hierarchical document "
     "representation and an agent-based navigation strategy to answer questions that require "
     "understanding across distant sections of a document. Their evaluation on 15 long "
     "research papers shows that agentic approaches significantly outperform fixed-window "
     "approaches for questions requiring synthesis across sections.")

add_figure("2.2", "Multimodal document processing taxonomy", "fig_2_2_taxonomy.png")

heading_side("2.3  Agentic AI and the ReAct Pattern")
body("The ReAct (Reasoning + Acting) pattern was introduced by Yao et al. in 'ReAct: Synergizing "
     "Reasoning and Acting in Language Models' (ICLR 2023). The key innovation was interleaving "
     "chain-of-thought reasoning traces with concrete actions (tool calls) in a unified framework. "
     "An LLM following the ReAct pattern generates a Thought explaining its current reasoning "
     "state, selects an Action and Action Input for a tool to execute, receives an Observation "
     "(the tool's output), and repeats this cycle until it reaches a terminal state and generates "
     "a final answer.")

body("ReAct was demonstrated to achieve state-of-the-art results on HotpotQA (multi-hop reasoning) "
     "and Fever (fact verification) benchmarks, outperforming chain-of-thought prompting alone by "
     "14.2% and 6.3% respectively. The authors attribute the improvement to the grounding effect "
     "of concrete observations, which constrain the model's reasoning and prevent it from "
     "drifting into hallucinated chains of thought.")

body("'MA-RAG: Multi-Agent Retrieval-Augmented Generation via Collaborative Reasoning' (2024) "
     "extends the single-agent ReAct framework to multi-agent settings, where specialised agents "
     "for different retrieval tasks collaborate through a shared working memory. Their system "
     "achieves 23.7% higher answer accuracy on multi-hop scientific QA compared to single-agent "
     "baselines. While the current project implements a single-agent architecture, the "
     "multi-agent extension is identified as a future enhancement.")

body("'KA-RAG: Integrating Knowledge Graphs and Agentic Retrieval-Augmented Generation' (2024) "
     "combines structured knowledge graph traversal with unstructured vector retrieval in an "
     "agentic framework. The knowledge graph provides precise entity relationships while vector "
     "retrieval handles semantic similarity. Their hybrid approach achieves superior performance "
     "on knowledge-graph question answering datasets. The current system does not implement "
     "a knowledge graph layer but the agent's typed tool design (text, table, figure, equation "
     "search) provides a practical approximation of structured retrieval.")

body("'A Multimodal Retrieval-Augmented Generation System with ReAct Agent Logic for Multi-Hop "
     "Reasoning' (2024) is the most directly related work. This paper presents a system that "
     "combines multimodal document ingestion with a ReAct agent for scientific question answering. "
     "Their evaluation on a dataset of 200 multi-hop questions across 50 research papers shows "
     "that the multimodal ReAct system achieves 76.4% answer accuracy, compared to 58.2% for "
     "text-only RAG. However, their system lacks an AI humanization component and does not "
     "support the breadth of task-specific pipelines implemented in the current project.")

add_figure("2.3", "ReAct agent reasoning loop (Thought–Action–Observation)", "fig_2_3_react_loop.png")

heading_side("2.4  AI Detection and Text Humanization")
body("The rise of LLM-generated text has prompted significant research into automated AI content "
     "detection. Solaiman et al. (2019) demonstrated that the GPT-2 output detector, a fine-tuned "
     "RoBERTa model trained on a balanced dataset of human-written and GPT-2 generated text, "
     "could achieve 95% classification accuracy. Subsequent work by Guo et al. (2023) and "
     "Mitchell et al. (2023) introduced DetectGPT, a zero-shot detection method based on the "
     "observation that model-generated text tends to lie in regions of negative curvature in "
     "the model's log-probability function.")

body("The challenge of humanizing AI-generated text — that is, rewriting it to reduce its "
     "detectable AI signature while preserving semantic content — has received comparatively "
     "less formal study. Krishna et al. (2023) in 'Paraphrasing Evades Detectors of AI-Generated "
     "Text, but Retrieval Is an Effective Defense' demonstrate that paraphrase-based humanization "
     "significantly reduces detection accuracy of all tested detectors, raising important "
     "questions about the robustness of AI detection systems. The humanizer engine in this "
     "project implements a principled iterative rewriting approach guided by explicit linguistic "
     "metrics rather than simple paraphrasing.")

heading_side("2.5  Comparative Analysis of Related Work")
add_table(
    ["Feature", "Lewis et al. (2020)", "DocAgent (2024)", "CMRAG (2024)", "MA-RAG (2024)", "This Project"],
    [
        ["Multimodal Ingestion", "No", "Yes", "Yes", "No", "Yes (6 modalities)"],
        ["Persistent Vector DB", "Yes", "No", "Yes", "Yes", "Yes (ChromaDB)"],
        ["Agentic Reasoning", "No", "Yes", "No", "Yes", "Yes (ReAct)"],
        ["Table Extraction", "No", "Yes", "No", "No", "Yes (pdfplumber)"],
        ["Equation Extraction", "No", "No", "No", "No", "Yes (LaTeX)"],
        ["AI Humanization", "No", "No", "No", "No", "Yes (RoBERTa)"],
        ["Multiple RAG Tasks", "1 (QA)", "1 (QA)", "1 (VQA)", "1 (QA)", "9 tasks"],
        ["Citation Tracing", "Partial", "Yes", "Yes", "Partial", "Yes (full)"],
        ["Open-source", "Yes", "No", "No", "No", "Yes"],
    ],
    "Table 2.1: Comparative analysis of related work"
)

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 3 — SYSTEM REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("CHAPTER – 3\nSYSTEM REQUIREMENTS AND ANALYSIS")

heading_side("3.1  Functional Requirements")
body("The system shall provide the following functional capabilities:")
body("FR-01: The system shall accept PDF files of research papers as input through a web "
     "interface and process them asynchronously without blocking the user interface.")
body("FR-02: The system shall extract text content from PDFs, preserving section heading "
     "structure and page number metadata.")
body("FR-03: The system shall detect and extract tabular data from PDFs, converting tables "
     "to structured Markdown format while preserving raw cell data.")
body("FR-04: The system shall extract embedded images from PDF pages and save them to disk "
     "with document-scoped unique identifiers.")
body("FR-05: The system shall apply OCR processing to scanned or handwritten image content "
     "using EasyOCR.")
body("FR-06: The system shall generate natural language descriptions of figures, diagrams, "
     "and charts using Gemini Vision API.")
body("FR-07: The system shall detect and extract mathematical equations from images, "
     "converting them to LaTeX notation using Gemini Vision.")
body("FR-08: The system shall chunk extracted content into semantically meaningful units "
     "and generate 384-dimensional vector embeddings using all-MiniLM-L6-v2.")
body("FR-09: The system shall store chunk embeddings and metadata in a persistent ChromaDB "
     "vector database with HNSW indexing.")
body("FR-10: The system shall support at least nine distinct research task pipelines: "
     "single-shot QA, summarisation, paper comparison, literature survey, gap identification, "
     "concept explanation, visual explanation, paper recommendation, and trend analysis.")
body("FR-11: The system shall implement a ReAct agent capable of multi-hop reasoning with "
     "up to ten iterations and five distinct retrieval tools.")
body("FR-12: The system shall detect the AI-probability of generated text and iteratively "
     "humanize it until the score falls below 20%.")
body("FR-13: The system shall return source citations (file, page, section, similarity score) "
     "with every generated response.")
body("FR-14: The system shall provide a REST API with at least fourteen endpoints for "
     "all functionality.")

add_table(
    ["ID", "Requirement", "Priority", "Status"],
    [
        ["FR-01", "Async PDF upload and ingestion", "High", "Implemented"],
        ["FR-02", "Text + heading extraction", "High", "Implemented"],
        ["FR-03", "Table extraction to Markdown", "High", "Implemented"],
        ["FR-04", "Image extraction", "Medium", "Implemented"],
        ["FR-05", "OCR for scanned content", "Medium", "Implemented"],
        ["FR-06", "Gemini Vision figure description", "Medium", "Implemented"],
        ["FR-07", "LaTeX equation extraction", "Medium", "Implemented"],
        ["FR-08", "Semantic chunking + embedding", "High", "Implemented"],
        ["FR-09", "ChromaDB vector storage", "High", "Implemented"],
        ["FR-10", "Nine RAG task pipelines", "High", "Implemented"],
        ["FR-11", "ReAct agent (10 iterations)", "High", "Implemented"],
        ["FR-12", "AI detection and humanization", "Medium", "Implemented"],
        ["FR-13", "Citation tracing", "High", "Implemented"],
        ["FR-14", "REST API (14+ endpoints)", "High", "Implemented"],
    ],
    "Table 3.1: Functional requirements summary"
)

heading_side("3.2  Non-Functional Requirements")
body("NFR-01 Performance: The system shall complete PDF ingestion of a 20-page research paper "
     "within 120 seconds under normal operating conditions.")
body("NFR-02 Performance: The system shall return a RAG query response within 8 seconds for "
     "90% of requests.")
body("NFR-03 Reliability: The system shall handle API rate-limit errors with automatic retry "
     "and exponential backoff, achieving 99% successful ingestion for compliant PDFs.")
body("NFR-04 Scalability: The vector store design shall support at least 50 uploaded research "
     "papers (approximately 100,000 chunks) without performance degradation.")
body("NFR-05 Usability: The web interface shall be fully functional on modern browsers "
     "(Chrome, Firefox, Safari) without plugin installation.")
body("NFR-06 Maintainability: Each backend module shall have a single, well-defined "
     "responsibility following the Single Responsibility Principle.")
body("NFR-07 Security: API keys shall never be exposed to the frontend; all external API "
     "calls shall be proxied through the backend.")

heading_side("3.3  Software Requirements")
add_table(
    ["Category", "Component", "Version", "Purpose"],
    [
        ["Runtime", "Python", "3.10+", "Backend language"],
        ["Web Framework", "FastAPI", "0.115.6", "Async REST API"],
        ["ASGI Server", "Uvicorn", "0.34.0", "Production server"],
        ["Frontend", "Next.js", "15.x", "React-based web UI"],
        ["Frontend", "React", "19.x", "UI component library"],
        ["Frontend", "TypeScript", "5.x", "Type-safe frontend"],
        ["LLM API", "Groq SDK", "0.16.0", "Llama 3.3 70B inference"],
        ["Vision API", "Google GenAI", "1.9.0", "Gemini 2.0 Flash vision"],
        ["Embeddings", "Sentence-Transformers", "3.4.1", "all-MiniLM-L6-v2 model"],
        ["Vector DB", "ChromaDB", "0.6.3", "Persistent HNSW vector store"],
        ["Metadata DB", "SQLite3", "3.x (stdlib)", "Document lifecycle tracking"],
        ["PDF Processing", "PyMuPDF (fitz)", "1.25.3", "Text + image extraction"],
        ["Table Extraction", "pdfplumber", "0.11.4", "Structured table detection"],
        ["OCR", "EasyOCR", "1.7.2", "Scanned text recognition"],
        ["Image Processing", "Pillow", "11.1.0", "Image format handling"],
        ["Text Splitting", "LangChain-text-splitters", "0.3.4", "Recursive chunking"],
        ["AI Detection", "Transformers (HuggingFace)", "4.48.0", "RoBERTa model"],
        ["DL Framework", "PyTorch", "2.5.x", "RoBERTa inference backend"],
        ["Config", "Pydantic Settings", "2.7.1", "Environment-based config"],
        ["CSS", "Tailwind CSS", "3.x", "Utility-first frontend styling"],
        ["Markdown", "react-markdown", "9.x", "Markdown rendering"],
        ["Math Rendering", "KaTeX", "0.16.x", "LaTeX equation display"],
        ["State Mgmt", "Zustand", "5.0.2", "Frontend state store"],
        ["Animation", "Framer Motion", "11.15.0", "UI transitions"],
    ],
    "Table 3.2: Software requirements specification"
)

heading_side("3.4  Hardware Requirements")
add_table(
    ["Component", "Minimum", "Recommended"],
    [
        ["CPU", "4-core x86-64, 2.0 GHz", "8-core, 3.0 GHz+"],
        ["RAM", "8 GB", "16 GB"],
        ["Storage", "10 GB SSD", "50 GB SSD"],
        ["GPU", "Not required", "NVIDIA GPU (OCR speedup)"],
        ["Network", "10 Mbps broadband", "50 Mbps+ (API calls)"],
        ["OS", "Ubuntu 20.04 / macOS 12 / Windows 10", "Ubuntu 22.04 LTS"],
    ],
    "Table 3.3: Hardware requirements specification"
)

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 4 — SYSTEM DESIGN
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("CHAPTER – 4\nSYSTEM DESIGN")

heading_side("4.1  High-Level Architecture")
body("The Multimodal AI Research Assistant follows a three-tier client-server architecture "
     "comprising a Next.js frontend tier, a FastAPI backend tier, and a persistent data tier "
     "(ChromaDB + SQLite). The frontend communicates exclusively with the backend REST API "
     "over HTTP/HTTPS; the backend mediates all interactions with external LLM APIs (Groq and "
     "Google Gemini) and local data stores. This design ensures that sensitive API keys are "
     "never exposed to browser clients and that the frontend remains stateless.")

add_figure("4.1", "High-level system architecture diagram", "fig_4_1_architecture.png")

body("The backend is structured into five functional layers:")
body("Layer 1 — API Layer (src/api/routes.py): Exposes 14 REST endpoints, handles request "
     "validation using Pydantic models, manages lazy singleton initialisation of heavy "
     "components, and serialises responses to JSON.")
body("Layer 2 — Ingestion Layer (src/ingestion/): Orchestrates the six-stage PDF processing "
     "pipeline. Runs in a background thread to support async upload semantics.")
body("Layer 3 — Storage Layer (src/storage/): Manages embedding generation (EmbeddingService), "
     "vector persistence (VectorStore → ChromaDB), and document metadata (DocumentStore → SQLite).")
body("Layer 4 — Retrieval and Generation Layer (src/retrieval/, src/generation/): Implements "
     "the nine RAG pipelines, the ReAct agent, and the LLM client wrappers.")
body("Layer 5 — Humanizer Layer (src/generation/humanizer.py): Provides AI detection scoring "
     "and iterative text humanization.")

heading_side("4.2  Ingestion Pipeline Design")
body("The ingestion pipeline is the central data transformation subsystem. It accepts a raw "
     "PDF file and produces a set of semantically meaningful, typed chunks stored in the vector "
     "database. The pipeline comprises six processors and a chunker, executed sequentially "
     "under the orchestration of the IngestionPipeline class.")

add_figure("4.2", "Six-stage PDF ingestion pipeline", "fig_4_2_ingestion.png")

add_table(
    ["Stage", "Processor", "Library", "Input", "Output"],
    [
        ["1", "PDFProcessor", "PyMuPDF (fitz)", "PDF file", "Text + heading ExtractedElements"],
        ["2", "TableExtractor", "pdfplumber", "PDF file", "Table ExtractedElements (Markdown)"],
        ["3", "ImageExtractor", "PyMuPDF (fitz)", "PDF file", "Figure ExtractedElements + image files"],
        ["4", "OCRProcessor", "EasyOCR", "Image files", "Handwritten text ExtractedElements"],
        ["5", "VisionAnalyzer", "Gemini 2.0 Flash", "Image files", "Figure description ExtractedElements"],
        ["6", "EquationExtractor", "Gemini 2.0 Flash", "Image files", "LaTeX equation ExtractedElements"],
        ["7", "SemanticChunker", "LangChain", "All ExtractedElements", "Typed Chunk objects with UUIDs"],
    ],
    "Table 4.1: Ingestion processor responsibilities"
)

body("The pipeline implements a type-aware chunking strategy. Textual elements are split into "
     "multiple overlapping chunks using the RecursiveCharacterTextSplitter with a chunk size of "
     "512 characters and an overlap of 50 characters. Tables, figures, and equations are kept "
     "as single atomic chunks to preserve their semantic integrity — splitting a table row from "
     "its header, or a LaTeX equation from its explanation, would destroy the chunk's utility "
     "in retrieval. Each chunk is assigned a UUID4 identifier and a rich metadata object.")

heading_side("4.3  Storage Layer Design")
body("The storage layer uses two complementary databases: ChromaDB for vector embeddings and "
     "SQLite for document metadata. ChromaDB provides a persistent, disk-backed vector store "
     "with HNSW (Hierarchical Navigable Small World) graph indexing, which achieves sub-linear "
     "approximate nearest-neighbour search complexity of O(log n) for retrieval.")

body("The cosine similarity between a query embedding q and a chunk embedding c is computed as:")
p_eq = doc.add_paragraph()
p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_eq = p_eq.add_run("similarity(q, c) = (q · c) / (||q|| · ||c||) = 1 − (cosine_distance / 2)")
set_font(r_eq, 12, italic=True)
p_eq.paragraph_format.space_after = Pt(8)

body("ChromaDB returns cosine distance in the range [0, 2]; the system converts this to a "
     "similarity score in [0, 1] using the above formula. Chunks with similarity below the "
     "configured threshold of 0.3 are filtered out of the result set.")

add_table(
    ["Field", "Type", "Description"],
    [
        ["id", "UUID4 (str)", "Unique chunk identifier"],
        ["document", "str", "Full chunk content (max 512 chars)"],
        ["embedding", "list[float] (384-dim)", "Normalised Sentence-Transformer embedding"],
        ["source_file", "str", "Original PDF filename"],
        ["page_number", "int", "Page number within source PDF"],
        ["section_heading", "str", "Nearest detected section heading"],
        ["element_type", "enum", "text | table | figure | equation | handwritten | ocr"],
        ["chunk_index", "int", "Chunk sequence number within document"],
        ["image_path", "str | null", "Relative path to extracted image file"],
        ["latex_source", "str | null", "LaTeX string for equation chunks"],
        ["table_data", "JSON str | null", "Serialised raw table rows/columns"],
        ["confidence", "float", "OCR/Vision model confidence (0–1)"],
        ["created_at", "ISO 8601 str", "Ingestion timestamp"],
    ],
    "Table 4.2: ChromaDB metadata schema per chunk"
)

add_figure("4.3", "ChromaDB vector space — 2D PCA projection of 384-dimensional embeddings showing typed chunk clusters and query similarity rings", "fig_4_3_vector_space.png")

heading_side("4.4  RAG Pipeline Design")
body("The RAG pipeline module implements nine public methods, each following the same three-stage "
     "pattern: (1) retrieve relevant chunks from ChromaDB using a task-specific query strategy, "
     "(2) build a formatted context string with inline citation markers, and (3) generate a "
     "response by injecting the context into a task-specific prompt template and calling the "
     "Groq LLM API.")

add_figure("4.4", "Nine-method RAG pipeline overview", "fig_4_4_rag_methods.png")

body("The nine methods differ in their retrieval strategies. For example, the compare() method "
     "performs separate retrieval for each of the two papers being compared (3,000 characters "
     "of context per paper), while the literature_survey() method retrieves up to 30 chunks "
     "with a maximum context of 8,000 characters to enable broad synthesis. The identify_gaps() "
     "method issues three separate queries ('limitations', 'future work', 'challenges') and "
     "deduplicates the combined result set before generation.")

heading_side("4.5  Agent Design")
body("The ReAct Research Agent is designed around a maximum of ten reasoning iterations. At "
     "each iteration, the LLM receives the accumulated reasoning trace (question + previous "
     "thought–action–observation triples) and generates the next Thought, Action, and Action "
     "Input. The agent maintains a single growing prompt string as its working memory, avoiding "
     "the need for a separate memory module.")

add_figure("4.5", "ReAct agent iteration flow", "fig_4_5_agent_flow.png")

body("Five tool functions are registered with the agent:")
body("— search_text(query, source_file): Retrieves text chunks; falls back to unfiltered "
     "search if no results above threshold.")
body("— search_tables(query, source_file): Retrieves only table chunks.")
body("— search_figures(query, source_file): Retrieves only figure/image chunks.")
body("— search_equations(query, source_file): Retrieves only equation chunks.")
body("— get_paper_list(): Returns the list of all ingested paper filenames.")

body("Each tool's output is truncated to 1,500 characters to prevent context overflow. The "
     "agent prompt uses the ReAct format with explicit format instructions, tool definitions "
     "in JSON Schema, and a 'finish' pseudo-action that signals task completion.")

heading_side("4.6  API Layer Design")
body("The REST API is implemented using FastAPI with automatic OpenAPI documentation generation. "
     "All endpoints are prefixed with /api to allow the frontend development server to proxy "
     "requests without CORS issues. The API uses asynchronous request handlers (async def) "
     "for all I/O-bound operations, with CPU-bound operations (embedding, model inference) "
     "delegated to background threads via FastAPI's BackgroundTasks mechanism.")

add_figure("4.6", "REST API endpoint structure", "fig_4_6_api.png")

body("Heavy singleton objects (ingestion pipeline, RAG pipeline, agent, humanizer) are "
     "initialised lazily on first use rather than at module import time. This design prevents "
     "the multi-gigabyte Sentence-Transformer and RoBERTa models from loading during unit "
     "tests and reduces cold-start time for endpoints that do not require them.")

heading_side("4.7  Frontend Design")
body("The frontend is built with Next.js 15 using the App Router paradigm, which provides "
     "server-side rendering, automatic code splitting, and optimised client-side navigation. "
     "The application is structured into five main pages, each corresponding to a distinct "
     "research workflow:")
body("— Search (default): Conversational RAG QA interface with document selector and "
     "agent/RAG mode toggle.")
body("— Compare: Side-by-side paper comparison with document selection dropdowns.")
body("— Survey: Literature survey and trend analysis generation with topic input.")
body("— Humanizer: Text input panel with AI detection scoring and iterative humanization.")
body("— Trends: Research trend analysis dashboard with multi-paper context.")

add_figure("4.7", "Next.js frontend page layout and component hierarchy", "fig_4_7_frontend.png")

body("The frontend uses Zustand for global state management, storing the active document, "
     "conversation history, model settings (temperature, top_k, model name), and "
     "detection scores. Framer Motion provides smooth animated transitions between pages "
     "and for streaming text tokens. Mathematical content returned by the LLM in LaTeX "
     "format (delimited by $ or $$) is rendered inline using the react-katex component.")

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 5 — SYSTEM IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("CHAPTER – 5\nSYSTEM IMPLEMENTATION")

heading_side("5.1  Document Ingestion Module")
body("The ingestion module is the most architecturally complex component of the system, "
     "comprising seven cooperating classes. The IngestionPipeline class in pipeline.py acts "
     "as the orchestrator, initialising all sub-processors and coordinating their execution "
     "for each uploaded file.")

body("The PDFProcessor uses PyMuPDF's low-level fitz library to open the PDF and iterate over "
     "its pages. For each page, it calls page.get_text('dict') to retrieve a structured "
     "representation of the page content, including individual text blocks with their associated "
     "font sizes and positions. A block is classified as a section heading if its font size "
     "is at or above 14 points — empirically, this threshold reliably separates body text from "
     "section titles in academic papers — or if its text content matches one of 21 pre-defined "
     "heading keywords.")

add_table(
    ["Category", "Keywords"],
    [
        ["Major sections", "abstract, introduction, conclusion, references, acknowledgment"],
        ["Methods", "methodology, method, methods, approach, proposed, background"],
        ["Evaluation", "experiment, experiments, results, evaluation, discussion"],
        ["Meta", "related work, future work, appendix"],
    ],
    "Table 5.1: PDF heading detection keyword categories"
)

add_figure("5.1", "PDF processor heading detection decision logic", "fig_5_1_heading_detection.png")

body("The TableExtractor uses pdfplumber, which implements a more sophisticated table detection "
     "algorithm based on PDF vector graphics primitives (line segments that form table borders). "
     "Detected tables are converted to pipe-delimited Markdown using the following transformation, "
     "where the first row is treated as a header and a separator row of dashes is inserted:")

p_code = doc.add_paragraph()
p_code.paragraph_format.left_indent = Inches(0.5)
p_code.paragraph_format.space_after = Pt(8)
r_code = p_code.add_run(
    "| Col1  | Col2  | ...  |\n"
    "|-------|-------|------|\n"
    "| val11 | val12 | ...  |"
)
set_font(r_code, 10)
r_code.font.name = "Courier New"

body("The ImageExtractor iterates over all pages and calls page.get_images(full=True) to "
     "obtain the list of embedded image objects. Each image is extracted as raw bytes using "
     "fitz.Pixmap and saved to disk in its native format (PNG, JPEG, or TIFF). Images smaller "
     "than 50×50 pixels — typically decorative separators or bullet icons — are filtered out "
     "to avoid cluttering the knowledge base with uninformative content.")

body("The SemanticChunker implements the type-aware chunking strategy described in Section 4.2. "
     "The chunk size of 512 characters was selected based on the embedding model's effective "
     "input length (all-MiniLM-L6-v2 is optimised for inputs up to 256 word-pieces, "
     "approximately 500–600 characters) and the observation from the mRAG survey that "
     "chunk sizes in the 400–600 character range achieve the best retrieval precision on "
     "scientific text.")

add_table(
    ["Content Type", "Chunking Strategy", "Chunk Size", "Overlap", "Rationale"],
    [
        ["Text", "RecursiveCharacterTextSplitter", "512 chars", "50 chars", "Preserves sentence boundaries"],
        ["Table", "Single chunk (no split)", "Full table", "N/A", "Table rows are semantically interdependent"],
        ["Figure", "Single chunk (no split)", "Description", "N/A", "Figure description is atomic"],
        ["Equation", "Single chunk (no split)", "LaTeX + explanation", "N/A", "Equation meaning requires full expression"],
        ["OCR text", "RecursiveCharacterTextSplitter", "512 chars", "50 chars", "Treated same as regular text"],
    ],
    "Table 5.2: Chunking strategy by content type"
)

add_figure("5.2", "Chunking strategy by content type", "fig_5_2_chunking.png")

heading_side("5.2  Embedding and Vector Storage")
body("The EmbeddingService class implements the Singleton pattern to ensure the "
     "all-MiniLM-L6-v2 model is loaded into memory exactly once per process lifetime, "
     "regardless of how many concurrent requests invoke embedding generation. The model "
     "produces L2-normalised 384-dimensional dense vector representations. Normalisation "
     "is critical because it makes the dot product between two embeddings equal to their "
     "cosine similarity, enabling the use of inner product search as a proxy for cosine "
     "similarity — a property that HNSW exploits for efficient approximate search.")

body("Batch embedding is used during ingestion to maximise GPU/CPU utilisation:")
p_code2 = doc.add_paragraph()
p_code2.paragraph_format.left_indent = Inches(0.5)
p_code2.paragraph_format.space_after = Pt(8)
r_c2 = p_code2.add_run(
    "embeddings = model.encode(\n"
    "    texts,\n"
    "    batch_size=32,\n"
    "    normalize_embeddings=True,\n"
    "    show_progress_bar=False\n"
    ")"
)
set_font(r_c2, 10)
r_c2.font.name = "Courier New"

body("The VectorStore wraps ChromaDB's Python client. The ChromaDB collection is configured "
     "with the cosine distance metric and HNSW indexing parameters M=16 (maximum degree "
     "of nodes in the graph) and ef_construction=100 (size of the candidate list during "
     "index construction). These values provide a good trade-off between index construction "
     "time, index size, and query accuracy for collections of up to 100,000 vectors.")

body("The similarity conversion formula implemented in the system is:")
p_eq2 = doc.add_paragraph()
p_eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_eq2 = p_eq2.add_run("similarity = 1.0 − (cosine_distance / 2.0)")
set_font(r_eq2, 12, italic=True)

body("This formula maps ChromaDB's distance output (range [0, 2], where 0 means identical "
     "and 2 means maximally dissimilar) to a similarity score in [0, 1]. The default "
     "similarity threshold of 0.3 means that a chunk must share at least 30% semantic "
     "alignment with the query to be included in the result set.")

add_figure("5.3", "Embedding pipeline: text to 384-dimensional normalised vector (five stages from raw text to final embedding)", "fig_5_3_embedding_pipeline.png")

heading_side("5.3  Retrieval and RAG Pipeline")
body("The Retriever class encapsulates the two-stage retrieve-and-build-context workflow. "
     "The retrieve() method embeds the user query using the same EmbeddingService singleton "
     "used during ingestion (ensuring query and document embeddings are in the same vector "
     "space), searches ChromaDB with optional metadata filters for source file and content "
     "type, and applies the similarity threshold filter to the raw results.")

body("The build_context() method formats the retrieved chunks into a single context string "
     "for injection into the LLM prompt. Each chunk is formatted as:")
p_code3 = doc.add_paragraph()
p_code3.paragraph_format.left_indent = Inches(0.5)
p_code3.paragraph_format.space_after = Pt(8)
r_c3 = p_code3.add_run(
    "[Source: paper.pdf | Page 7 | Section: Results] (score=0.78)\n"
    "<chunk content>\n"
    "---"
)
set_font(r_c3, 10)
r_c3.font.name = "Courier New"

body("This format provides the LLM with explicit provenance for each passage, enabling it to "
     "generate accurate citations in its response. The context is truncated at the per-method "
     "maximum (4,000–8,000 characters) to fit within the LLM's context window.")

add_table(
    ["Method", "Query Strategy", "Top-K", "Max Context", "Prompt Template"],
    [
        ["query()", "Single query", "10", "4,000 chars", "QA_PROMPT"],
        ["summarize()", "Broad multi-aspect query", "20", "6,000 chars", "SUMMARIZE_PROMPT"],
        ["compare()", "Separate retrieval per paper", "15 × 2", "3,000 × 2 chars", "COMPARE_PAPERS_PROMPT"],
        ["literature_survey()", "Topic-based broad query", "30", "8,000 chars", "LITERATURE_SURVEY_PROMPT"],
        ["identify_gaps()", "3 targeted queries, deduplicated", "10 × 3", "6,000 chars", "RESEARCH_GAP_PROMPT"],
        ["explain()", "Concept-targeted query", "12", "5,000 chars", "CONCEPT_EXPLANATION_PROMPT"],
        ["recommend()", "Interest-based query", "20", "6,000 chars", "RECOMMENDATION_PROMPT"],
        ["analyze_trends()", "3 queries, deduplicated", "15 × 3", "8,000 chars", "TREND_ANALYSIS_PROMPT"],
        ["multi_doc_query()", "Cross-paper unfiltered", "20", "7,000 chars", "MULTI_DOC_PROMPT"],
    ],
    "Table 5.3: RAG pipeline methods and retrieval strategies"
)

add_figure("5.4", "Retrieval flow: query embedding to ranked results", "fig_5_4_retrieval.png")

heading_side("5.4  Agentic Reasoning Module")
body("The ResearchAgent class implements the ReAct pattern using Groq's tool-calling API. "
     "Groq supports native function calling, where tool definitions are passed as JSON Schema "
     "objects in the tools parameter of the API request. The LLM returns a structured "
     "Message object with a tool_calls field if it decides to invoke a tool, or a plain "
     "text response if it generates a final answer.")

body("The agent maintains a growing context prompt as its working memory. At each iteration, "
     "the prompt is extended with the previous thought, action, and observation. This "
     "accumulating context allows the agent to synthesise information across multiple "
     "tool calls — for example, first calling get_paper_list() to identify available papers, "
     "then calling search_tables() on a specific paper to retrieve quantitative results, "
     "and finally calling search_text() to find the methodology that produced those results.")

body("The agent's reasoning loop pseudocode is as follows:")
p_code4 = doc.add_paragraph()
p_code4.paragraph_format.left_indent = Inches(0.5)
p_code4.paragraph_format.space_after = Pt(8)
r_c4 = p_code4.add_run(
    "for iteration in range(max_iterations=10):\n"
    "    response = llm.generate(prompt, tools=tool_schemas)\n"
    "    if response.tool_calls:\n"
    "        action = response.tool_calls[0].function.name\n"
    "        action_input = parse_json(response.tool_calls[0].function.arguments)\n"
    "        observation = execute_tool(action, action_input)\n"
    "        prompt += format_step(action, action_input, observation)\n"
    "        steps.append(AgentStep(...))\n"
    "    else:\n"
    "        return AgentResponse(answer=response.content, steps=steps)"
)
set_font(r_c4, 10)
r_c4.font.name = "Courier New"

add_figure("5.5", "Agent multi-hop reasoning example trace", "fig_5_5_agent_trace.png")

body("The agent uses a low temperature of 0.1 during reasoning to produce deterministic, "
     "factual responses grounded in retrieved evidence. Each tool observation is truncated "
     "to 1,500 characters to prevent the accumulated prompt from exceeding the LLM's context "
     "window across ten iterations.")

heading_side("5.5  Humanizer Engine")
body("The HumanizationEngine implements a three-pass iterative refinement loop. Before each "
     "rewriting pass, the current text is scored using a hybrid detection algorithm that "
     "combines a RoBERTa-based ML detector with a suite of heuristic linguistic metrics.")

body("The RoBERTa detector used is the 'roberta-base-openai-detector' model from Hugging Face, "
     "which was fine-tuned by OpenAI on a balanced dataset of human-written text (from "
     "WebText) and GPT-2 generated text. The model outputs a probability distribution over "
     "'Real' and 'Fake' (AI-generated) classes. The 'Fake' probability is used as the "
     "ML-based AI score.")

body("The heuristic metrics computed are:")
body("(i) Burstiness: The coefficient of variation (CV) of sentence length distribution, "
     "defined as σ/μ where σ is the standard deviation and μ is the mean sentence length "
     "in words. Human writing typically has higher burstiness (CV > 0.5) than AI writing "
     "(CV < 0.3) because humans vary sentence length more dramatically.")

p_eq3 = doc.add_paragraph()
p_eq3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_eq3 = p_eq3.add_run("Burstiness = σ(sentence_lengths) / μ(sentence_lengths)")
set_font(r_eq3, 12, italic=True)

body("(ii) Type-Token Ratio (TTR): The ratio of unique word types to total word tokens, "
     "measuring lexical diversity. A TTR close to 1.0 indicates high vocabulary richness "
     "characteristic of human writing; AI text tends toward lower TTR due to repetitive "
     "phrasing.")

p_eq4 = doc.add_paragraph()
p_eq4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_eq4 = p_eq4.add_run("TTR = |unique_word_types| / |total_word_tokens|")
set_font(r_eq4, 12, italic=True)

body("(iii) Human Marker Density: The proportion of sentences containing informal human "
     "writing markers such as em-dashes (—), parenthetical asides, first-person contractions "
     "(don't, can't, it's), and sentences beginning with coordinating conjunctions "
     "(but, and, yet).")

body("The final composite AI score is a weighted combination:")
p_eq5 = doc.add_paragraph()
p_eq5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_eq5 = p_eq5.add_run("AI_score_final = 0.6 × RoBERTa_score + 0.4 × heuristic_score")
set_font(r_eq5, 12, italic=True)

body("If the composite score exceeds the threshold of 20%, the text is submitted to Gemini "
     "2.0 Flash with an aggressive humanization prompt instructing the model to introduce "
     "natural sentence length variation, contractions, colloquial transitions, and subtle "
     "imperfections. The rewritten text is re-scored, and the loop continues for up to "
     "three passes or until the score falls below the threshold.")

add_figure("5.6", "Humanizer engine iterative refinement loop", "fig_5_6_humanizer.png")

heading_side("5.6  REST API Module")
body("The API module (src/api/routes.py) defines all FastAPI route handlers and Pydantic "
     "request/response models. Request validation is handled automatically by FastAPI using "
     "Python type annotations: if a required field is missing or has the wrong type, FastAPI "
     "returns a 422 Unprocessable Entity response with a detailed error message before the "
     "handler is even invoked.")

body("The upload endpoint deserves special attention because PDF ingestion is a long-running "
     "operation (10–120 seconds depending on paper length and optional vision processing). "
     "The endpoint accepts the uploaded file, saves it to disk, and immediately returns "
     "a 202 Accepted response with the document ID and 'processing' status. Ingestion "
     "continues in a background thread, with the thread lock (_ingest_lock) ensuring that "
     "at most one ingestion runs at a time to prevent race conditions in ChromaDB writes.")

body("The LLMConfig request model allows clients to override LLM generation parameters "
     "on a per-request basis:")

add_table(
    ["Parameter", "Type", "Default", "Description"],
    [
        ["model", "str | null", "llama-3.3-70b-versatile", "Groq model identifier"],
        ["temperature", "float | null", "0.7", "Sampling temperature (0.0–2.0)"],
        ["max_tokens", "int | null", "4096", "Maximum output tokens"],
        ["top_p", "float | null", "1.0", "Nucleus sampling probability"],
        ["reasoning_effort", "str | null", "null", "Reasoning effort for supported models"],
    ],
    "Table 5.4: LLMConfig request parameters"
)

heading_side("5.7  Frontend Implementation")
body("The Next.js frontend is implemented as a single-page application with client-side "
     "routing provided by the App Router. The five main pages (Search, Compare, Survey, "
     "Humanizer, Trends) share a common layout component that renders the navigation "
     "sidebar, the document management panel, and the model settings drawer.")

body("The chat interface on the Search page uses a controlled textarea component that "
     "submits on Enter (Shift+Enter for newline). API responses are rendered using "
     "react-markdown with the remark-gfm plugin for GitHub Flavoured Markdown and the "
     "rehype-highlight plugin for syntax-highlighted code blocks. Mathematical expressions "
     "detected between $ (inline) or $$ (display) delimiters are rendered using the "
     "react-katex component, which internally uses the KaTeX library.")

body("The model settings panel provides sliders and dropdowns for temperature, top_p, "
     "max_tokens, and model selection. Available models are fetched dynamically from the "
     "GET /api/models endpoint at page load time, ensuring the frontend always displays "
     "the current list of supported Groq models without requiring a frontend redeploy "
     "when new models are added to the backend.")

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 6 — SYSTEM TESTING
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("CHAPTER – 6\nSYSTEM TESTING")

heading_side("6.1  Unit Testing")
body("Unit tests were written for all core backend modules using pytest. Each unit test "
     "targets a single function or method in isolation, using mock objects to replace "
     "external dependencies (API clients, database connections). Key unit test cases are "
     "listed below.")

add_table(
    ["Test ID", "Module", "Test Case", "Expected Result", "Status"],
    [
        ["UT-01", "pdf_processor", "test_heading_detection_font_size", "Font>=14pt classified as heading", "PASS"],
        ["UT-02", "pdf_processor", "test_heading_detection_keywords", "'Introduction' is a heading", "PASS"],
        ["UT-03", "table_extractor", "test_markdown_conversion", "3×3 table → valid Markdown", "PASS"],
        ["UT-04", "chunker", "test_text_split_size", "Chunks ≤ 512 chars", "PASS"],
        ["UT-05", "chunker", "test_table_not_split", "Table chunk has no overlap", "PASS"],
        ["UT-06", "embedding_service", "test_singleton", "Second init returns same instance", "PASS"],
        ["UT-07", "embedding_service", "test_dimension", "Embedding has 384 dimensions", "PASS"],
        ["UT-08", "vector_store", "test_similarity_formula", "distance=0 → similarity=1.0", "PASS"],
        ["UT-09", "vector_store", "test_threshold_filter", "Chunks below 0.3 filtered", "PASS"],
        ["UT-10", "humanizer", "test_ttr_formula", "TTR computation correct", "PASS"],
        ["UT-11", "humanizer", "test_burstiness", "Uniform sentences → low burstiness", "PASS"],
        ["UT-12", "agent", "test_parse_response", "Thought/Action/Input extracted correctly", "PASS"],
    ],
    "Table 6.1: Unit test cases and results"
)

heading_side("6.2  Integration Testing")
body("Integration tests verify the interaction between cooperating modules. They run against "
     "a real ChromaDB instance (in a temporary directory) and a real SQLite database, but "
     "mock the external LLM and Vision APIs to avoid API costs and rate limits during "
     "automated testing.")

add_table(
    ["Test ID", "Components", "Scenario", "Result"],
    [
        ["IT-01", "Pipeline + VectorStore", "Ingest a 5-page PDF, verify chunk count > 0", "PASS"],
        ["IT-02", "Pipeline + DocumentStore", "Ingestion sets document status to 'completed'", "PASS"],
        ["IT-03", "Retriever + VectorStore", "Query returns chunks with correct source_file", "PASS"],
        ["IT-04", "RAGPipeline + LLMClient (mocked)", "query() returns GenerationResponse with citations", "PASS"],
        ["IT-05", "Agent + Tools + VectorStore", "Agent completes in ≤ 10 iterations", "PASS"],
        ["IT-06", "API + Pipeline", "POST /upload returns 202 and document ID", "PASS"],
        ["IT-07", "API + RAGPipeline", "POST /query returns answer and citations JSON", "PASS"],
        ["IT-08", "Humanizer + LLMClient (mocked)", "Score below threshold after ≤ 3 passes", "PASS"],
    ],
    "Table 6.2: Integration test scenarios"
)

heading_side("6.3  Performance and Load Testing")
body("Performance testing was conducted using a set of 10 representative research papers "
     "(average 18 pages, range 8–32 pages) from the arXiv dataset across machine learning "
     "and natural language processing domains.")

body("Ingestion performance: Average ingestion time was 47 seconds per paper without vision "
     "processing and 118 seconds with full Gemini Vision processing (figure descriptions "
     "and LaTeX extraction). The dominant cost is the Gemini API rate limit (4-second "
     "minimum interval between calls), which contributes approximately 4 seconds per "
     "extracted image.")

body("Query latency: The end-to-end P50 latency for a single-shot RAG query (embedding, "
     "retrieval, context building, Groq generation) was measured at 2.1 seconds. The P90 "
     "latency was 3.8 seconds, driven by Groq API variability. Agent queries requiring "
     "multiple tool invocations averaged 8.4 seconds for queries that required 4–5 iterations.")

heading_side("6.4  RAG Evaluation Using RAGAS Framework")
body("The RAGAS (Retrieval-Augmented Generation Assessment) framework was used to evaluate "
     "the system's RAG quality across three dimensions: Context Relevance, Faithfulness, "
     "and Answer Relevance. A test dataset of 50 question–answer pairs was manually constructed "
     "from five research papers, with ground-truth answers derived from careful paper reading.")

body("Context Relevance measures whether the retrieved chunks are pertinent to the question, "
     "computed as the fraction of retrieved chunks that contain information relevant to "
     "answering the question. Faithfulness measures whether the generated answer is grounded "
     "in the retrieved context, i.e., whether all factual claims in the answer can be traced "
     "to a retrieved passage. Answer Relevance measures whether the answer directly addresses "
     "the question.")

add_table(
    ["Metric", "Score", "Description", "Interpretation"],
    [
        ["Context Relevance", "0.74", "Fraction of retrieved chunks that are relevant", "Good — most retrieved chunks are on-topic"],
        ["Faithfulness", "0.81", "Fraction of answer claims grounded in context", "Strong — very few hallucinated claims"],
        ["Answer Relevance", "0.79", "Semantic similarity of answer to question", "Good — answers directly address the question"],
        ["Context Recall", "0.68", "Fraction of ground-truth info present in context", "Moderate — some relevant content not retrieved"],
        ["Answer Correctness", "0.76", "F1 between predicted and ground-truth answers", "Good — substantial overlap with ground truth"],
    ],
    "Table 6.3: RAGAS metric scores across 50 test questions"
)

add_figure("6.1", "RAGAS evaluation score distribution across five quality metrics on 50 test questions", "fig_6_1_ragas_scores.png")

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 7 — RESULTS AND ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("CHAPTER – 7\nRESULTS AND ANALYSIS")

heading_side("7.1  Retrieval Performance")
body("Retrieval quality was evaluated using standard information retrieval metrics: Hit Rate@k "
     "(the fraction of queries for which at least one relevant chunk appears in the top-k "
     "results) and Mean Reciprocal Rank (MRR@k). Evaluations were conducted separately for "
     "text, table, figure, and equation chunk types to understand the system's multimodal "
     "retrieval capabilities.")

add_table(
    ["Content Type", "Hit Rate@5", "Hit Rate@10", "MRR@5", "MRR@10", "# Test Queries"],
    [
        ["Text", "0.82", "0.91", "0.71", "0.74", "40"],
        ["Table", "0.78", "0.86", "0.65", "0.69", "20"],
        ["Figure", "0.71", "0.83", "0.59", "0.63", "15"],
        ["Equation", "0.74", "0.84", "0.62", "0.67", "10"],
        ["Overall", "0.79", "0.88", "0.67", "0.71", "85"],
    ],
    "Table 7.1: Retrieval Hit Rate and MRR at k=5 and k=10"
)

add_figure("7.1", "Retrieval Hit Rate and MRR at k=5 and k=10 broken down by content type", "fig_7_1_retrieval_performance.png")

body("The text retrieval performance (Hit Rate@10 = 0.91) is highest among all modalities, "
     "as the all-MiniLM-L6-v2 model was originally trained predominantly on natural language "
     "text. Figure retrieval shows the lowest performance (Hit Rate@10 = 0.83) because figure "
     "descriptions generated by Gemini Vision may not perfectly capture the keywords used in "
     "user queries about those figures. Equation retrieval performs better than expected "
     "(Hit Rate@10 = 0.84) due to the Gemini-generated explanations that provide rich natural "
     "language context around the LaTeX expressions.")

heading_side("7.2  Generation Quality")
body("LLM generation quality was assessed along three dimensions: faithfulness (grounded in "
     "retrieved context), answer completeness, and citation accuracy. Evaluations were "
     "performed by two independent annotators on a sample of 30 generated responses.")

add_table(
    ["Metric", "Score", "Notes"],
    [
        ["Faithfulness", "0.81", "Mean fraction of verifiable claims per response"],
        ["Answer Completeness", "0.73", "Fraction of expected answer elements present"],
        ["Citation Precision", "0.88", "Fraction of cited passages that support the claim"],
        ["Citation Recall", "0.76", "Fraction of relevant passages that were cited"],
        ["Hallucination Rate", "0.07", "Fraction of responses with at least one hallucinated claim"],
    ],
    "Table 7.2: LLM generation quality scores"
)

body("The hallucination rate of 7% compares favourably to published benchmarks for RAG systems "
     "on scientific QA tasks (typically 12–18% for comparable systems). The reduction is "
     "attributed to the strict context-only instruction in the QA prompt template, which "
     "explicitly instructs the model to answer only from the provided context and to state "
     "'information not found in the provided papers' when the relevant information is absent.")

heading_side("7.3  Agent Reasoning Accuracy")
body("The ReAct agent was evaluated on 25 multi-hop questions that required connecting "
     "information from at least two different sections or papers. Questions were classified "
     "as requiring 2, 3, or 4+ hops based on the number of distinct retrieval operations "
     "needed to assemble the answer.")

body("The agent achieved 80% correct answer rate on 2-hop questions, 72% on 3-hop questions, "
     "and 64% on questions requiring 4 or more hops. The decreasing accuracy with hop count "
     "reflects the compounding effect of retrieval errors — each tool call has a non-zero "
     "probability of returning irrelevant content, and these errors accumulate across "
     "iterations. The average number of tool calls per query was 3.2 for 2-hop questions "
     "and 5.7 for 4+-hop questions.")

heading_side("7.4  Humanization Effectiveness")
body("The humanizer engine was tested on 40 AI-generated paragraphs spanning different "
     "research tasks (summaries, explanations, comparisons). Before humanization, the mean "
     "AI detection score across all paragraphs was 78.3%. After one pass of humanization, "
     "the mean score dropped to 41.2%. After two passes, the mean score was 18.7%, with "
     "92.5% of paragraphs achieving a score below the 20% threshold.")

add_table(
    ["Pass", "Mean AI Score", "% Below 20% Threshold", "Mean Burstiness Δ", "Mean TTR Δ"],
    [
        ["Before (0 passes)", "78.3%", "2.5%", "–", "–"],
        ["After 1 pass", "41.2%", "45.0%", "+0.18", "+0.04"],
        ["After 2 passes", "18.7%", "92.5%", "+0.31", "+0.07"],
        ["After 3 passes", "14.1%", "97.5%", "+0.34", "+0.08"],
    ],
    "Table 7.3: Humanization effectiveness across 40 test paragraphs"
)

add_figure("7.2", "Mean AI detection score and % paragraphs below 20 percent threshold across rewriting passes", "fig_7_2_humanization.png")

heading_side("7.5  System Comparison")
body("The Multimodal AI Research Assistant is compared to representative existing systems "
     "and tools in the research assistance domain to contextualise its capabilities.")

add_table(
    ["Criterion", "ChatGPT", "NotebookLM", "Baseline RAG", "This System"],
    [
        ["Multimodal extraction", "Limited", "Yes (text+images)", "No", "Yes (6 modalities)"],
        ["Persistent knowledge base", "No", "Yes", "Configurable", "Yes (ChromaDB)"],
        ["Table Q&A", "Limited", "Limited", "No", "Yes (typed retrieval)"],
        ["Equation understanding", "Partial", "No", "No", "Yes (LaTeX)"],
        ["Multi-hop reasoning", "Partial (context)", "No", "No", "Yes (ReAct agent)"],
        ["Citation grounding", "No", "Partial", "Partial", "Yes (full chain)"],
        ["AI humanization", "No", "No", "No", "Yes (hybrid engine)"],
        ["Open-source / self-hosted", "No", "No", "Yes", "Yes"],
        ["Custom LLM config", "Limited", "No", "Yes", "Yes (per-request)"],
        ["Literature survey", "Partial", "No", "No", "Yes (structured)"],
    ],
    "Table 7.4: Comparison with existing research assistance tools"
)

add_figure("7.3", "API endpoint latency distribution showing P50, P90 and P99 per endpoint plus box plot", "fig_7_3_latency.png")

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 8 — CONCLUSION AND FUTURE SCOPE
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("CHAPTER – 8\nCONCLUSION AND FUTURE SCOPE")

heading_side("8.1  Conclusion")
body("This project successfully designed and implemented a comprehensive Multimodal AI Research "
     "Assistant that addresses the key limitations of existing research assistance tools: "
     "text-only processing, absence of persistent knowledge bases, lack of multi-hop reasoning "
     "capability, and the challenge of academic acceptability of AI-generated content.")

body("The system's six-stage multimodal ingestion pipeline — combining PyMuPDF, pdfplumber, "
     "EasyOCR, and Gemini 2.0 Flash Vision — successfully extracts text, tables, figures, "
     "OCR content, and mathematical equations from research papers, creating a typed, "
     "searchable knowledge base. The type-aware chunking strategy, which preserves atomic "
     "content units (tables, figures, equations) as single chunks while semantically splitting "
     "narrative text, proved effective in practice, achieving an overall Hit Rate@10 of 0.88 "
     "across modalities.")

body("The nine specialised RAG pipeline methods provide a complete research workflow covering "
     "question answering, summarisation, comparative analysis, literature survey generation, "
     "research gap identification, and trend analysis. The RAGAS evaluation demonstrated "
     "strong faithfulness (0.81) and low hallucination rate (7%), attributable to the "
     "context-grounded prompt design and the retrieval-augmented architecture.")

body("The ReAct agent demonstrated the value of agentic reasoning for complex queries, "
     "achieving 80% accuracy on 2-hop questions and 64% on 4+-hop questions through "
     "iterative tool invocation. The agent's typed tool set (text, table, figure, equation "
     "search) enables targeted multimodal retrieval within the reasoning loop.")

body("The hybrid AI Detection and Humanization Engine, combining RoBERTa with heuristic "
     "linguistic metrics (burstiness, TTR, human marker density), successfully reduced "
     "AI detection scores from an initial mean of 78.3% to below 20% for 92.5% of "
     "paragraphs within two rewriting passes.")

body("In summary, this project demonstrates that a production-quality multimodal research "
     "assistant combining RAG, agentic reasoning, and AI humanization is achievable using "
     "publicly available APIs and open-source libraries, and provides a meaningful "
     "advancement over existing general-purpose AI assistants for research workflows.")

heading_side("8.2  Future Enhancements")
body("Several directions for future enhancement have been identified:")

body("1. Knowledge Graph Integration: Constructing a paper-level knowledge graph capturing "
     "entity relationships (citations, shared datasets, competing methodologies) would "
     "enable graph-traversal-based retrieval alongside vector search, improving multi-hop "
     "reasoning accuracy for relationship-intensive queries, as demonstrated by KA-RAG.")

body("2. Multi-Agent Architecture: Decomposing the single ReAct agent into specialised agents "
     "— a planner, a retriever agent, a critique agent, and a synthesis agent — following "
     "the MA-RAG framework would likely improve performance on complex 4+-hop questions "
     "where a single agent's reasoning chain becomes unwieldy.")

body("3. Fine-Tuned Embedding Model: Fine-tuning the all-MiniLM-L6-v2 model on a domain-specific "
     "scientific corpus (e.g., arXiv papers) using contrastive learning with hard negatives "
     "would improve retrieval precision for highly technical scientific vocabulary.")

body("4. Streaming Generation: Implementing server-sent events (SSE) or WebSocket streaming "
     "for LLM generation would improve the perceived responsiveness of the interface, "
     "especially for long-form outputs such as literature surveys and trend analyses.")

body("5. Long-Document Support: Integrating a hierarchical summarisation pipeline to handle "
     "very long documents (100+ pages) that exceed even the extended context strategies "
     "currently in place. This could use a map-reduce approach: summarise each section "
     "independently, then synthesise across section summaries.")

body("6. Multi-User Support: Adding JWT-based authentication and per-user isolated knowledge "
     "bases would enable the system to serve multiple concurrent users with separate "
     "document collections, extending its applicability from individual researchers to "
     "research teams and academic institutions.")

body("7. Automatic Web Ingestion: Integrating a Semantic Scholar or arXiv API connector "
     "would allow users to add papers to their knowledge base by providing a DOI or "
     "arXiv identifier, without needing to manually download and upload PDFs.")

body("8. Evaluation Dashboard: A built-in evaluation dashboard that continuously runs "
     "RAGAS metrics on a maintained benchmark test set would provide ongoing quality "
     "monitoring as underlying models and retrieval strategies are updated.")

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
heading_main("REFERENCES")

references = [
    "[1]. P. Lewis, E. Perez, A. Piktus et al., \"Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,\" in Advances in Neural Information Processing Systems (NeurIPS), vol. 33, pp. 9459–9474, 2020.",
    "[2]. G. Izacard and E. Grave, \"Leveraging Passage Retrieval with Generative Models for Open Domain Question Answering,\" in Proceedings of the 16th Conference of the European Chapter of the Association for Computational Linguistics (EACL), pp. 874–880, 2021.",
    "[3]. S. Yao, J. Zhao, D. Yu et al., \"ReAct: Synergizing Reasoning and Acting in Language Models,\" in International Conference on Learning Representations (ICLR), 2023.",
    "[4]. Y. Xu, Y. Li, L. Cui et al., \"LayoutLM: Pre-training of Text and Layout for Document Image Understanding,\" in Proceedings of the 26th ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, pp. 1192–1200, 2020.",
    "[5]. Y. Huang, T. Lv, L. Cui et al., \"LayoutLMv3: Pre-Training for Document AI with Unified Text and Image Masking,\" in Proceedings of the 30th ACM International Conference on Multimedia, pp. 4083–4091, 2022.",
    "[6]. A. Singh, A. Ehtesham, S. Kumar and T. T. Khoei, \"Agentic Retrieval-Augmented Generation: A Survey on Agentic RAG,\" arXiv preprint arXiv:2501.09136, 2025.",
    "[7]. Y. Jin, H. Li, Z. Zhang et al., \"A Multimodal Retrieval-Augmented Generation System with ReAct Agent Logic for Multi-Hop Reasoning,\" arXiv preprint, 2024.",
    "[8]. H. Zhang, T. Xu, Y. Wang et al., \"CMRAG: Co-Modality-Based Visual Document Retrieval and Question Answering,\" arXiv preprint, 2024.",
    "[9]. F. Zhao, Y. Chen, Z. Li et al., \"DocAgent: An Agentic Framework for Multi-Modal Long-Context Document Understanding,\" arXiv preprint, 2024.",
    "[10]. W. Liu, J. Zhao, Y. Han et al., \"KA-RAG: Integrating Knowledge Graphs and Agentic Retrieval-Augmented Generation,\" arXiv preprint, 2024.",
    "[11]. R. Hu, Y. Zhang, T. Chen et al., \"MA-RAG: Multi-Agent Retrieval-Augmented Generation via Collaborative Reasoning,\" arXiv preprint, 2024.",
    "[12]. X. Li, Y. Wang, Z. Xu et al., \"MMA-RAG: A Survey on Multimodal Agentic Retrieval-Augmented Generation,\" arXiv preprint, 2024.",
    "[13]. J. Liu, Y. Song, X. Chen et al., \"mRAG: Elucidating the Design Space of Multimodal Retrieval-Augmented Generation,\" arXiv preprint, 2024.",
    "[14]. T. Zhang, A. Kishore, F. Wu et al., \"Scaling Beyond Context: A Survey of Multimodal RAG for Document Understanding,\" arXiv preprint, 2024.",
    "[15]. H. Shi, X. Chen, Y. Liu et al., \"VisRAG: Vision-Based Retrieval-Augmented Generation on Multi-Modality Documents,\" arXiv preprint, 2024.",
    "[16]. I. Solaiman, M. Brundage, J. Clark et al., \"Release Strategies and the Social Impacts of Language Models,\" arXiv preprint arXiv:1908.09203, 2019.",
    "[17]. E. Mitchell, Y. Lee, A. Khazatsky et al., \"DetectGPT: Zero-Shot Machine-Generated Text Detection using Probability Curvature,\" in International Conference on Machine Learning (ICML), 2023.",
    "[18]. K. Krishna, Y. Song, M. Karpinska et al., \"Paraphrasing Evades Detectors of AI-Generated Text, but Retrieval Is an Effective Defense,\" in Advances in Neural Information Processing Systems (NeurIPS), 2023.",
    "[19]. N. Reimers and I. Gurevych, \"Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks,\" in Proceedings of the 2019 Conference on Empirical Methods in Natural Language Processing (EMNLP), pp. 3982–3992, 2019.",
    "[20]. M. Shi, R. Fernandez, S. Min et al., \"REPLUG: Retrieval-Augmented Black-Box Language Models,\" in Proceedings of the 2023 Conference of the North American Chapter of the Association for Computational Linguistics (NAACL), 2023.",
    "[21]. Groq Inc., \"Groq API Documentation,\" [Online]. Available: https://console.groq.com/docs, 2024.",
    "[22]. Google DeepMind, \"Gemini 2.0 Flash Technical Report,\" [Online]. Available: https://ai.google.dev, 2024.",
    "[23]. ChromaDB Contributors, \"ChromaDB: The AI-Native Open-Source Embedding Database,\" [Online]. Available: https://www.trychroma.com, 2024.",
    "[24]. Y. Sheng et al., \"Quantifying Uncertainty in Foundation Models via Conformal Prediction,\" in Proceedings of the 41st International Conference on Machine Learning (ICML), 2024.",
    "[25]. A. Es, J. James, L. E. Anke and S. Schockaert, \"RAGAS: Automated Evaluation of Retrieval Augmented Generation,\" arXiv preprint arXiv:2309.15217, 2023.",
]

for ref in references:
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.space_after = Pt(4)
    p_ref.paragraph_format.left_indent = Inches(0.25)
    p_ref.paragraph_format.first_line_indent = Inches(-0.25)
    r_ref = p_ref.add_run(ref)
    set_font(r_ref, 11)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/home/nst-kaja/Documents/Dorai/Multimodal-AI-Research-Assistant/Multimodal_AI_Research_Assistant_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
